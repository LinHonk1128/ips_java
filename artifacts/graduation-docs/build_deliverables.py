from __future__ import annotations

import shutil
from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt


ROOT = Path(__file__).resolve().parent
TASK_TEMPLATE = ROOT / "任务书模板.docx"
OPENING_SOURCE = ROOT / "开题报告.docx"
TASK_OUTPUT = ROOT / "金宁丰-毕业设计任务书.docx"
OPENING_OUTPUT = ROOT / "金宁丰-开题报告-与论文一致版.docx"
REPORT_OUTPUT = ROOT / "开题报告与论文一致性检查说明.docx"


def set_run_font(run, size=10.5, bold=False, font_name="宋体"):
    run.font.name = font_name
    run.font.size = Pt(size)
    run.font.bold = bold
    r_pr = run._element.get_or_add_rPr()
    r_fonts = r_pr.get_or_add_rFonts()
    r_fonts.set(qn("w:eastAsia"), font_name)
    r_fonts.set(qn("w:ascii"), "Times New Roman")
    r_fonts.set(qn("w:hAnsi"), "Times New Roman")


def set_cell_margins(cell, top=60, start=90, bottom=60, end=90):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin_name, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin_name}"))
        if node is None:
            node = OxmlElement(f"w:{margin_name}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def clear_cell(cell):
    cell._tc.clear_content()


def add_cell_lines(cell, lines, size=10.0, line_spacing=1.05, first_bold=False):
    clear_cell(cell)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
    set_cell_margins(cell)
    for index, line in enumerate(lines):
        paragraph = cell.add_paragraph()
        paragraph.paragraph_format.space_before = Pt(0)
        paragraph.paragraph_format.space_after = Pt(0)
        paragraph.paragraph_format.line_spacing = line_spacing
        paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        run = paragraph.add_run(line)
        set_run_font(run, size=size, bold=(first_bold and index == 0))


def add_labeled_content(cell, heading, items, size=9.5):
    clear_cell(cell)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
    set_cell_margins(cell, top=45, start=80, bottom=45, end=80)
    p = cell.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.0
    run = p.add_run(heading)
    set_run_font(run, size=size, bold=True)
    for item in items:
        p = cell.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.line_spacing = 1.0
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        run = p.add_run(item)
        set_run_font(run, size=size)


def build_task_book():
    shutil.copy2(TASK_TEMPLATE, TASK_OUTPUT)
    doc = Document(TASK_OUTPUT)

    info = doc.paragraphs[1]
    info.clear()
    info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    info.paragraph_format.space_before = Pt(0)
    info.paragraph_format.space_after = Pt(4)
    run = info.add_run(
        "专业班级  软件222    学号  202201050631    姓名  金宁丰    下发日期  2026年4月13日"
    )
    set_run_font(run, size=11)

    table = doc.tables[0]
    add_cell_lines(
        table.rows[0].cells[1],
        ["基于RAG与知识画像的考研学习辅助系统设计与实现"],
        size=11,
    )
    table.rows[0].cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_cell_lines(
        table.rows[1].cells[1],
        ["个人知识库问答、知识画像与学习闭环设计"],
        size=10.5,
    )
    table.rows[1].cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    main_cell = table.rows[2].cells[1]
    clear_cell(main_cell)
    main_cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
    set_cell_margins(main_cell, top=45, start=80, bottom=45, end=80)
    main_sections = [
        (
            "主要内容：",
            [
                "1. 分析考研自学中资料分散、知识定位效率低、复习反馈弱、错题与知识点脱节及学习计划难以动态调整等问题，完成系统需求分析、总体架构设计、功能模块设计和数据库设计。",
                "2. 以知识片段（KnowledgeChunk）为核心知识单元，实现多类型资料的文本抽取、人工校正、语义切片与向量化处理，构建个人知识库。",
                "3. 实现基于个人知识库的智能问答，采用关键词检索与向量检索相结合的混合检索策略，通过RRF融合和规则重排序筛选知识片段，并返回资料来源引用。",
                "4. 根据知识片段的引用、用户反馈、错题关联和练习结果构建知识画像，分析知识覆盖率、掌握程度、薄弱知识点、遗忘风险和复习优先级。",
                "5. 实现错题复盘、知识片段关联、练习结果回写及个性化学习计划生成，形成“资料整理-知识检索-错题复盘-薄弱点识别-学习计划调整”的学习闭环。",
            ],
        ),
        (
            "要求：",
            [
                "1. 系统功能应与论文描述保持一致，完成资料管理、个人知识库问答、知识画像、错题复盘和学习计划等核心模块，并保证用户数据隔离。",
                "2. 完成主要业务流程、异常输入和访问权限测试，以及性能、兼容性、易用性和安全性等非功能测试。",
                "3. 按学校毕业设计规范完成论文撰写、系统图表整理、格式审查和答辩材料准备，做到结构完整、表述准确、术语统一。",
            ],
        ),
        (
            "参考文献：",
            [
                "[1] LEWIS P, PEREZ E, PIKTUS A, et al. Retrieval-augmented generation for knowledge-intensive NLP tasks[C]//NeurIPS. 2020:9459-9474.",
                "[2] 王士进,吴金泽,张浩天,等.可信的端到端深度学生知识画像建模方法[J].计算机研究与发展,2023,60(8):1822-1833.",
                "[3] 刘雪颖,云静,李博,等.基于大型语言模型的检索增强生成综述[J].计算机工程与应用,2025,61(13):1-25.",
                "[4] CORMACK G V, CLARKE C L A, BÜTTCHER S. Reciprocal rank fusion outperforms Condorcet and individual rank learning methods[C]//SIGIR. 2009:758-759.",
            ],
        ),
    ]
    for heading, items in main_sections:
        paragraph = main_cell.add_paragraph()
        paragraph.paragraph_format.space_before = Pt(0)
        paragraph.paragraph_format.space_after = Pt(0)
        paragraph.paragraph_format.line_spacing = 1.0
        run = paragraph.add_run(heading)
        set_run_font(run, size=8.3, bold=True)
        for item in items:
            paragraph = main_cell.add_paragraph()
            paragraph.paragraph_format.space_before = Pt(0)
            paragraph.paragraph_format.space_after = Pt(0)
            paragraph.paragraph_format.line_spacing = 1.0
            paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            run = paragraph.add_run(item)
            set_run_font(run, size=8.3)

    tech_items = [
        "1. 系统采用前后端分离架构。后端使用JDK 21、Spring Boot 3.3.5、Spring MVC、Spring Security、JWT和Spring Data JPA；前端使用Vue 3、Vite 5和ECharts。",
        "2. MySQL保存用户、学习档案、资料、知识片段、学习行为、错题和学习计划等结构化数据，本地uploads目录保存原始资料及附件；所有业务数据与当前用户绑定。",
        "3. PDFBox、Apache POI和Tesseract OCR分别用于PDF、Word和图片资料的文本抽取；知识片段采用“语义边界优先、长度阈值辅助”的切片策略，切片版本为2，最小长度约300字符、目标长度约800字符、最大长度约1100字符。",
        "4. Elasticsearch、Embedding服务、大语言模型和OCR服务作为增强能力接入。混合检索中关键词检索与向量检索分别召回20个候选片段，经RRF融合保留前20个结果，本地重排后最多选择5个知识片段构造回答上下文；外部服务异常时保留基础检索或本地统计能力。",
        "5. 系统应保证常用页面和核心接口正常运行，未登录访问、无效身份及跨用户资源访问应被拦截；测试覆盖主要业务流程、异常输入及访问权限。",
    ]
    add_labeled_content(table.rows[3].cells[1], "主要技术参数：", tech_items, size=8.6)

    schedule_items = [
        "第一阶段（4月13日-4月26日）：完善系统并开展联调，完成RAG检索流程端到端验证，修复核心功能缺陷，设计功能测试与异常场景测试用例。",
        "第二阶段（4月27日-5月11日）：完成系统部署说明和用户操作说明，撰写论文初稿，整理系统架构图、流程图、数据库ER图及系统页面等材料。",
        "第三阶段（5月12日-5月18日）：根据指导教师意见修改论文，补充系统测试章节与非功能测试结果，完善摘要、结论和参考文献。",
        "第四阶段（5月19日-5月30日）：结合评阅意见优化系统功能和论文内容，制作答辩PPT，准备系统演示环境并进行答辩演练。",
        "第五阶段（5月31日-6月6日）：完成论文查重、格式审查和最终排版，整理毕业设计相关材料。",
        "第六阶段（6月7日-6月11日）：参加毕业设计答辩，根据答辩意见修改论文，提交最终版论文和电子材料。",
    ]
    add_labeled_content(table.rows[4].cells[1], "进度及完成日期：", schedule_items, size=8.8)

    ref_cell = doc.tables[4].cell(0, 0)
    clear_cell(ref_cell)
    ref_cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
    set_cell_margins(ref_cell, top=100, start=160, bottom=80, end=160)
    title_p = ref_cell.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_p.paragraph_format.space_after = Pt(4)
    title_run = title_p.add_run("参    考    文    献")
    set_run_font(title_run, size=12, bold=True)

    doc.save(TASK_OUTPUT)


def all_paragraphs(doc):
    for paragraph in doc.paragraphs:
        yield paragraph
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    yield paragraph


def replace_paragraph_text(paragraph, old, new):
    if old not in paragraph.text:
        return False
    original_runs = list(paragraph.runs)
    font_name = None
    font_size = None
    bold = None
    if original_runs:
        font_name = original_runs[0].font.name
        font_size = original_runs[0].font.size
        bold = original_runs[0].font.bold
    text = paragraph.text.replace(old, new)
    paragraph.clear()
    run = paragraph.add_run(text)
    if font_name:
        run.font.name = font_name
    if font_size:
        run.font.size = font_size
    run.font.bold = bold
    r_pr = run._element.get_or_add_rPr()
    r_pr.get_or_add_rFonts().set(qn("w:eastAsia"), font_name or "宋体")
    return True


def revise_opening_report():
    shutil.copy2(OPENING_SOURCE, OPENING_OUTPUT)
    doc = Document(OPENING_OUTPUT)
    replacements = [
        (
            "形成基于SpringBoot、Vue3、MySQL、Elasticsearch和大语言模型接口的模块化实现方案",
            "形成基于Spring Boot、Vue 3和MySQL的前后端分离实现方案，并将Elasticsearch、Embedding、OCR和大语言模型接口作为可选增强能力接入",
        ),
        (
            "结合标题、句子边界和语义连续性进行知识切片，将长篇资料转化为可检索、可引用、可记录学习事件的知识切片。",
            "结合标题、句子边界和语义连续性进行语义切片，将长篇资料转化为可检索、可引用、可记录学习事件的知识片段。",
        ),
        (
            "验证大模型回答能够准确根据个人资料生成并提供来源引用。",
            "验证系统能够基于个人资料生成回答、返回来源引用，并在外部服务不可用时执行降级处理。",
        ),
        (
            "拟采用Spring Boot构建后端服务，Vue 3构建前端界面，MySQL存储业务数据，Elasticsearch提供全文和向量检索；按照模块逐步实现并通过联调持续修正设计。",
            "拟采用Spring Boot构建后端服务，Vue 3构建前端界面，MySQL存储业务数据，Elasticsearch作为可选增强能力提供全文和向量检索；按照模块逐步实现并通过联调持续修正设计，外部服务不可用时保留基础检索和本地统计能力。",
        ),
        (
            "（4）对比与实验验证法。针对关键词检索、向量检索和混合检索设计典型问题集，比较召回结果；针对资料处理、来源引用、画像更新、错题回写和计划生成设计功能测试与异常场景测试。",
            "（4）系统测试法。围绕用户访问控制、资料处理、个人知识库问答、来源反馈回写、错题复盘、知识画像更新和学习计划转化设计功能测试与异常场景测试，并对性能、兼容性、易用性和安全性进行非功能测试。",
        ),
        (
            "修复核心功能缺陷，编写单元测试与集成测试用例。交付可运行的系统原型及测试文档。",
            "修复核心功能缺陷，设计功能测试与异常场景测试用例。交付可运行的系统原型及测试记录。",
        ),
        (
            "补充系统测试章节与性能评估数据",
            "补充系统测试章节与非功能测试结果",
        ),
    ]
    counts = {}
    for old, new in replacements:
        count = 0
        for paragraph in all_paragraphs(doc):
            if replace_paragraph_text(paragraph, old, new):
                count += 1
        counts[old[:24]] = count
    doc.save(OPENING_OUTPUT)
    print("Opening report replacements:", counts)


def set_doc_defaults(doc):
    style = doc.styles["Normal"]
    style.font.name = "宋体"
    style.font.size = Pt(11)
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    for section in doc.sections:
        section.top_margin = section.bottom_margin = Pt(56.7)
        section.left_margin = section.right_margin = Pt(68)


def add_heading(doc, text, level=1):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    set_run_font(run, size=14 if level == 1 else 12, bold=True, font_name="黑体")
    return p


def add_body(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Pt(22)
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.space_after = Pt(0)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    run = p.add_run(text)
    set_run_font(run, size=11)
    return p


def build_consistency_report():
    doc = Document()
    set_doc_defaults(doc)
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(12)
    run = title.add_run("开题报告与毕业论文一致性检查说明")
    set_run_font(run, size=18, bold=True, font_name="黑体")

    add_body(
        doc,
        "检查以《基于RAG与知识画像的考研学习辅助系统设计与实现》论文初稿为优先依据，重点核对题目、研究对象、技术路线、功能模块、测试范围和进度安排。修正版开题报告仅调整与论文实际内容存在冲突或容易造成过度表述的内容，未改变原模板结构。",
    )

    add_heading(doc, "一、已修正的冲突", 1)
    items = [
        "1. 核心术语：开题报告将“知识切片”同时用于处理过程和数据对象，论文则以“知识片段（KnowledgeChunk）”为核心知识单元。修正版保留“语义切片”作为处理方法，将处理结果统一称为“知识片段”。",
        "2. 技术架构：开题报告原文容易将Elasticsearch理解为不可缺少的基础组件；论文将MySQL和本地存储作为基础数据支撑，将Elasticsearch、Embedding、OCR和大语言模型接口作为可选增强能力，并设置降级流程。修正版已按论文调整。",
        "3. 问答验证目标：开题报告原文使用“验证回答准确”的表述，但论文主要完成检索、来源返回、反馈回写和降级流程的功能验证，尚未使用标准问题集量化回答准确率。修正版已改为验证回答生成、来源引用和降级处理的功能正确性。",
        "4. 研究方法：开题报告原计划对关键词检索、向量检索和混合检索开展典型问题集对比；论文明确将标准问题集、召回率和回答准确性的量化比较列为后续工作。修正版已改为系统功能测试和非功能测试。",
        "5. 测试安排：开题报告进度中写有“单元测试与集成测试”，论文实际以功能测试、连续业务场景测试和非功能测试为主，并将自动化接口测试列为后续工作。修正版已删除已完成单元/集成测试的暗示。",
        "6. 性能表述：开题报告原文强调“性能评估数据”，论文完成了非功能测试，但仍认为自动化测试和性能测试存在完善空间。修正版使用“非功能测试结果”，避免超过论文结论。",
    ]
    for item in items:
        add_body(doc, item)

    add_heading(doc, "二、与论文保持一致的内容", 1)
    aligned = [
        "题目、研究场景和问题定义一致，均围绕考研自学中的资料分散、问答来源不清、复习反馈弱、错题与知识点脱节及计划调整缺少依据展开。",
        "系统核心业务模块一致：资料分类与上传管理、基于个人知识库的智能问答、知识画像分析、错题复盘与知识点关联、个性化学习计划生成；用户学习档案与基础支撑作为运行前置条件。",
        "技术路线一致：Spring Boot、Vue 3和MySQL构成前后端分离基础架构，结合文档解析、OCR、Embedding、Elasticsearch、RRF和大语言模型实现增强能力。",
        "学习闭环一致，统一表述为“资料整理-知识检索-错题复盘-薄弱点识别-学习计划调整”。",
    ]
    for index, item in enumerate(aligned, 1):
        add_body(doc, f"{index}. {item}")

    add_heading(doc, "三、论文自身仍建议统一的表述", 1)
    add_body(
        doc,
        "论文内部仍存在“Spring Boot/SpringBoot/SpringBoot3”“Vue 3/Vue3”等书写形式混用，以及“实现基于从个人知识库的问答服务”“知识画像分错题复盘”“主要包括要包括”“系统根据系统根据”等明显文字问题。本次未修改论文原文件；任务书和修正版开题报告采用规范写法，并以论文实际实现内容为准。",
    )

    doc.save(REPORT_OUTPUT)


def main():
    build_task_book()
    revise_opening_report()
    build_consistency_report()
    for path in (TASK_OUTPUT, OPENING_OUTPUT, REPORT_OUTPUT):
        print(path.name, path.stat().st_size)


if __name__ == "__main__":
    main()
