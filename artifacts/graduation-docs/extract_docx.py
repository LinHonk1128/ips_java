from __future__ import annotations

import json
import sys
import zipfile
from pathlib import Path

from docx import Document
from lxml import etree


W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
NS = {"w": W_NS}


def xml_text(element: etree._Element) -> str:
    parts = element.xpath(".//w:t/text() | .//w:tab/text()", namespaces=NS)
    return "".join(parts).strip()


def extract_docx(path: Path) -> dict:
    doc = Document(path)
    result = {
        "file": str(path),
        "paragraphs": [],
        "tables": [],
        "sections": [],
        "headers": [],
        "footers": [],
        "comments": [],
        "tracked_changes": {"insertions": 0, "deletions": 0},
    }

    for index, paragraph in enumerate(doc.paragraphs):
        text = paragraph.text.strip()
        if text:
            result["paragraphs"].append(
                {
                    "index": index,
                    "style": paragraph.style.name if paragraph.style else "",
                    "text": text,
                }
            )

    for table_index, table in enumerate(doc.tables):
        rows = []
        for row in table.rows:
            rows.append([cell.text.strip() for cell in row.cells])
        result["tables"].append({"index": table_index, "rows": rows})

    for index, section in enumerate(doc.sections):
        result["sections"].append(
            {
                "index": index,
                "page_width": section.page_width,
                "page_height": section.page_height,
                "top_margin": section.top_margin,
                "bottom_margin": section.bottom_margin,
                "left_margin": section.left_margin,
                "right_margin": section.right_margin,
            }
        )
        header_text = "\n".join(p.text.strip() for p in section.header.paragraphs if p.text.strip())
        footer_text = "\n".join(p.text.strip() for p in section.footer.paragraphs if p.text.strip())
        if header_text:
            result["headers"].append({"section": index, "text": header_text})
        if footer_text:
            result["footers"].append({"section": index, "text": footer_text})

    with zipfile.ZipFile(path) as archive:
        names = set(archive.namelist())
        document_root = etree.fromstring(archive.read("word/document.xml"))
        result["tracked_changes"]["insertions"] = len(document_root.xpath(".//w:ins", namespaces=NS))
        result["tracked_changes"]["deletions"] = len(document_root.xpath(".//w:del", namespaces=NS))

        if "word/comments.xml" in names:
            comments_root = etree.fromstring(archive.read("word/comments.xml"))
            for comment in comments_root.xpath(".//w:comment", namespaces=NS):
                result["comments"].append(
                    {
                        "id": comment.get(f"{{{W_NS}}}id"),
                        "author": comment.get(f"{{{W_NS}}}author", ""),
                        "date": comment.get(f"{{{W_NS}}}date", ""),
                        "text": xml_text(comment),
                    }
                )

        result["all_document_text"] = [
            xml_text(paragraph)
            for paragraph in document_root.xpath(".//w:p", namespaces=NS)
            if xml_text(paragraph)
        ]

    return result


def write_text_dump(data: dict, output: Path) -> None:
    lines = [f"# {Path(data['file']).name}", "", "## Paragraphs"]
    for item in data["paragraphs"]:
        lines.append(f"[P{item['index']}|{item['style']}] {item['text']}")

    lines.extend(["", "## Tables"])
    for table in data["tables"]:
        lines.append(f"### Table {table['index']}")
        for row_index, row in enumerate(table["rows"]):
            lines.append(f"[R{row_index}] " + " | ".join(row))

    lines.extend(["", "## Headers"])
    for item in data["headers"]:
        lines.append(f"[Section {item['section']}] {item['text']}")

    lines.extend(["", "## Footers"])
    for item in data["footers"]:
        lines.append(f"[Section {item['section']}] {item['text']}")

    lines.extend(["", "## Comments"])
    for item in data["comments"]:
        lines.append(f"[{item['id']}] {item['author']}: {item['text']}")

    lines.extend(["", "## All document XML text"])
    for index, text in enumerate(data["all_document_text"]):
        lines.append(f"[X{index}] {text}")

    output.write_text("\n".join(lines), encoding="utf-8")


def main() -> None:
    for filename in sys.argv[1:]:
        path = Path(filename).resolve()
        data = extract_docx(path)
        json_path = path.with_suffix(".extracted.json")
        text_path = path.with_suffix(".extracted.txt")
        json_path.write_text(json.dumps(data, ensure_ascii=False, indent=2), encoding="utf-8")
        write_text_dump(data, text_path)
        print(f"{path.name}: {len(data['paragraphs'])} paragraphs, "
              f"{len(data['tables'])} tables, {len(data['comments'])} comments")


if __name__ == "__main__":
    main()
