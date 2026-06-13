from pathlib import Path
from shutil import copy2

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt


ROOT = Path(__file__).resolve().parents[1]
SOURCE = ROOT / "docs" / "论文前六章内容整合稿.docx"
OUTPUT = ROOT / "docs" / "论文前七章内容整合稿.docx"


def set_run_font(run, size=None, bold=None):
    run.font.name = "Times New Roman"
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), "宋体")
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shading = tc_pr.find(qn("w:shd"))
    if shading is None:
        shading = OxmlElement("w:shd")
        tc_pr.append(shading)
    shading.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=100, bottom=80, end=100):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def format_table(table, widths):
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    for row_index, row in enumerate(table.rows):
        for column_index, cell in enumerate(row.cells):
            cell.width = widths[column_index]
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            set_cell_margins(cell)
            if row_index == 0:
                set_cell_shading(cell, "D9EAF7")
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_before = Pt(0)
                paragraph.paragraph_format.space_after = Pt(0)
                paragraph.paragraph_format.line_spacing = 1.0
                paragraph.alignment = (
                    WD_ALIGN_PARAGRAPH.CENTER
                    if column_index in (0, len(widths) - 1)
                    else WD_ALIGN_PARAGRAPH.LEFT
                )
                for run in paragraph.runs:
                    set_run_font(run, 9.5, row_index == 0)


def add_caption(document, text):
    paragraph = document.add_paragraph(style="Caption")
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run(text)
    set_run_font(run, 9)


def add_body(document, text):
    paragraph = document.add_paragraph(text, style="Normal")
    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    paragraph.paragraph_format.first_line_indent = Pt(21)
    return paragraph


def add_heading(document, text, level):
    paragraph = document.add_heading(text, level=level)
    if level == 1:
        paragraph.paragraph_format.page_break_before = True
    return paragraph


def replace_paragraph_text(paragraph, text):
    paragraph.text = text
    for run in paragraph.runs:
        set_run_font(run, 10)


def update_chapter_overview(document):
    chapter_five = next(
        paragraph
        for paragraph in document.paragraphs
        if paragraph.text.startswith("第 5 章为系统实现")
    )
    chapter_six = next(
        paragraph
        for paragraph in document.paragraphs
        if paragraph.text.startswith("第 6 章为系统测试与总结")
    )
    chapter_seven = next(
        paragraph
        for paragraph in document.paragraphs
        if paragraph.text.startswith("第7章为总结与展望")
    )

    replace_paragraph_text(
        chapter_five,
        "第 5 章为数据库设计。主要介绍数据库概念结构、逻辑结构和物理结构，说明用户、"
        "资料、知识片段、学习事件、错题和学习计划等数据之间的关系。",
    )
    replace_paragraph_text(
        chapter_six,
        "第 6 章为系统实现。主要说明用户学习档案、资料上传与知识切片、知识问答、"
        "知识画像、错题复盘和个性化学习计划等核心模块的实现过程。",
    )
    inserted = chapter_seven.insert_paragraph_before(
        "第 7 章为系统测试。主要对系统核心功能和关键业务流程进行测试，并对测试结果进行分析。"
    )
    inserted.style = chapter_seven.style
    for run in inserted.runs:
        set_run_font(run, 10)
    replace_paragraph_text(
        chapter_seven,
        "第 8 章为总结与展望。主要总结本文完成的工作，分析系统当前不足，并说明后续优化方向。",
    )


def main():
    copy2(SOURCE, OUTPUT)
    document = Document(OUTPUT)
    update_chapter_overview(document)

    add_heading(document, "系统测试", 1)
    add_body(
        document,
        "本章选取系统核心业务链路进行功能测试，重点检查用户访问控制、个人资料库问答、"
        "学习反馈回写、错题练习和学习计划转化是否符合预期。测试以演示数据和接口操作为主，"
        "不再重复说明各模块的实现机制。",
    )

    add_heading(document, "7.1 测试环境", 2)
    add_body(
        document,
        "测试在本地开发环境中完成。为验证外部服务不可用时的基本可用性，接口测试关闭"
        "Elasticsearch，并清空模型与向量服务密钥，知识问答使用系统本地检索与摘要结果。",
    )
    environment_rows = [
        ("测试项", "环境配置"),
        ("操作系统", "Windows 11"),
        ("后端环境", "JDK 21、Spring Boot 3.3.5"),
        ("前端环境", "Vue 3、Vite 5"),
        ("数据环境", "MySQL 8、demo 演示数据"),
        ("测试方式", "项目构建、REST 接口与核心业务操作"),
    ]
    table = document.add_table(rows=len(environment_rows), cols=2)
    for row_index, values in enumerate(environment_rows):
        for column_index, value in enumerate(values):
            table.cell(row_index, column_index).text = value
    format_table(table, [Cm(3.6), Cm(11.4)])
    add_caption(document, "表7-1 系统测试环境")

    add_heading(document, "7.2 核心功能测试", 2)
    add_body(
        document,
        "根据前文确定的系统重点，本节不对全部页面逐项测试，而是围绕“身份验证—知识获取—"
        "学习反馈—复习行动”主链路设计测试用例。测试结果如表7-2所示。",
    )
    test_rows = [
        ("编号", "测试内容", "操作与预期结果", "实际结果", "结论"),
        (
            "T01",
            "访问控制",
            "未携带令牌访问资料接口，应拒绝请求",
            "返回403，受保护接口未泄露业务数据",
            "通过",
        ),
        (
            "T02",
            "用户登录",
            "使用演示账号登录，应返回JWT并允许访问个人数据",
            "登录成功，令牌可用于后续接口请求",
            "通过",
        ),
        (
            "T03",
            "知识库读取",
            "读取科目文件夹和画像概览，应返回已构建的知识数据",
            "返回15个文件夹、67个知识片段，覆盖率为88.06%",
            "通过",
        ),
        (
            "T04",
            "资料库问答",
            "围绕指定文件夹提问，应返回答案及可追溯来源",
            "返回529字答案和2条来源，来源含片段编号与摘要",
            "通过",
        ),
        (
            "T05",
            "来源反馈",
            "对问答来源提交“忘记了”，应更新片段学习统计",
            "反馈成功，目标片段错误反馈次数更新",
            "通过",
        ),
        (
            "T06",
            "错题练习回写",
            "提交一次答题结果，应同步更新关联知识片段",
            "练习记录保存成功，1个关联片段完成回写",
            "通过",
        ),
        (
            "T07",
            "画像建议转计划",
            "将复习建议加入计划，应生成可查询的学习任务",
            "任务创建成功，来源标记为AI，并可正常删除",
            "通过",
        ),
        (
            "T08",
            "项目构建",
            "分别执行后端打包与前端生产构建，应无编译错误",
            "后端打包成功，前端构建成功",
            "通过",
        ),
    ]
    table = document.add_table(rows=len(test_rows), cols=5)
    for row_index, values in enumerate(test_rows):
        for column_index, value in enumerate(values):
            table.cell(row_index, column_index).text = value
    format_table(table, [Cm(1.1), Cm(2.5), Cm(5.0), Cm(5.0), Cm(1.4)])
    table.rows[0]._tr.get_or_add_trPr().append(OxmlElement("w:tblHeader"))
    add_caption(document, "表7-2 核心功能测试结果")

    add_heading(document, "7.3 测试结果分析", 2)
    add_body(
        document,
        "表7-2中的测试用例均达到预期结果。未登录请求能够被鉴权模块拦截，登录后的资料、"
        "画像、错题和计划数据可以正常读取，说明基础访问控制和个人业务数据链路可用。知识库"
        "问答在关闭外部检索与模型服务后仍能返回答案和来源，表明系统的本地降级流程能够支持"
        "基本查询。",
    )
    add_body(
        document,
        "反馈测试表明，问答来源反馈和错题练习结果能够继续写入知识片段统计；画像建议也能够"
        "转化为带来源标记的学习任务。因此，系统重点实现的“知识获取—反馈记录—状态分析—"
        "计划调整”闭环可以正常运行。当前测试主要验证功能正确性，尚未对高并发访问、长期运行"
        "稳定性以及不同大模型的回答质量进行量化比较，这些内容可在后续工作中补充。",
    )

    add_heading(document, "7.4 本章小结", 2)
    add_body(
        document,
        "本章对系统核心功能进行了简要测试。测试结果显示，用户鉴权、知识库读取与问答、"
        "学习反馈回写、错题练习和学习计划转化均能按预期完成，前后端项目也能够正常构建。"
        "系统已具备支撑个人考研资料管理和学习反馈闭环的基本运行能力。",
    )

    document.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    main()
