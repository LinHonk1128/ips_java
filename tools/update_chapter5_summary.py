from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement


ROOT = Path(__file__).resolve().parents[1]
SOURCE = ROOT / "docs" / "论文前五章内容整合稿.docx"
OUTPUT = ROOT / "docs" / "论文前五章内容整合稿-补充5.4.docx"


def insert_paragraph_after(paragraph, text, style=None):
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    new_paragraph = paragraph._parent.add_paragraph()
    new_paragraph._p.getparent().remove(new_paragraph._p)
    new_p.getparent().replace(new_p, new_paragraph._p)

    if style is not None:
        new_paragraph.style = style
    if text:
        new_paragraph.add_run(text)
    return new_paragraph


def main():
    doc = Document(SOURCE)
    paragraphs = doc.paragraphs

    chapter_53 = next(p for p in paragraphs if p.text.strip() == "5.3 数据库物理结构设计")
    chapter_53.style = doc.styles["Heading 2"]

    number_fixes = {
        "用户学习档案表：": "(2) 用户学习档案表：",
        "学习资料文件夹表：": "(3) 学习资料文件夹表：",
        "知识片段表：": "(5) 知识片段表：",
        "错题表：": "(8) 错题表：",
    }
    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        for prefix, replacement in number_fixes.items():
            if text.startswith(prefix):
                paragraph.text = replacement + text[len(prefix) :]
                break

    old_summary = next(
        p
        for p in doc.paragraphs
        if p.text.strip().startswith("以上数据表中，知识片段表、知识片段事件表")
    )
    old_summary._element.getparent().remove(old_summary._element)

    heading = next(p for p in doc.paragraphs if p.text.strip() == "5.4本章小结")
    heading.text = "5.4 本章小结"
    heading.style = doc.styles["Heading 2"]

    summary_paragraphs = [
        (
            "本章依据系统业务需求，依次完成了数据库概念设计、逻辑结构设计和物理结构设计。"
            "概念设计明确了用户、学习档案、资料文件夹、学习资料、知识片段、行为事件、错题和学习计划等主要实体；"
            "逻辑结构设计进一步说明了实体间的层级关系与多对多关联；物理结构设计则将上述实体落实为可供系统持久化使用的数据表。"
            "整体上，数据库以用户编号确定数据归属，以知识片段作为连接资料内容和学习行为的核心数据单元，在保证不同用户数据隔离的同时，为各业务模块提供统一的数据基础。"
        ),
        (
            "上述设计并非只用于保存静态资料，而是承担学习闭环中的数据传递和状态沉淀。"
            "用户上传的资料经文件夹和资料文件表组织后被切分为知识片段，知识片段为智能问答、来源引用和定制练题提供内容依据；"
            "知识片段事件表持续记录引用、理解反馈和练习结果，错题与知识片段关联表将错题暴露的问题定位到具体知识内容，错题练习事件表则保留复盘结果及其时间信息。"
            "这些数据经过学习画像模块聚合后，可用于识别掌握程度较低或长期未复习的知识点，并进一步生成复习建议或学习计划。"
            "用户后续的问答、练习和错题复盘又会产生新的行为记录，使数据库中的学习状态随实际学习过程持续更新，从而实现“资料入库—知识使用—学习反馈—薄弱点识别—计划调整”的闭环。"
        ),
        (
            "从系统实现角度看，分层且关联明确的数据结构减少了资料、错题和学习状态之间的重复存储，主键与关联字段为跨模块查询、结果回写和权限校验提供了稳定依据；"
            "知识片段表中的汇总字段能够支持掌握度和复习优先级的快速计算，事件表保留的明细记录又能满足趋势分析和后续规则调整的需要。"
            "因此，本章的数据库设计不仅完成了各类业务数据的持久化，还为智能问答的来源追溯、学习画像的动态计算、错题复盘的知识关联以及学习计划的个性化生成提供了实现支撑。"
        ),
    ]

    current = heading
    for text in summary_paragraphs:
        current = insert_paragraph_after(current, text, style=doc.styles["Normal"])

    for paragraph in doc.paragraphs:
        if not paragraph.text.strip() and paragraph._p.getprevious() is not None:
            previous_text = "".join(paragraph._p.getprevious().itertext()).strip()
            if previous_text == "5.4 本章小结":
                paragraph._element.getparent().remove(paragraph._element)

    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    main()
