from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


OUT = r"D:\Project\biyesheji\ips-java\毕业设计开题报告-202201050631-金宁丰.docx"


def set_run_font(run, size=12, bold=False, name="宋体"):
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    run._element.rPr.rFonts.set(qn("w:ascii"), "Times New Roman")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Times New Roman")
    run.font.size = Pt(size)
    run.bold = bold


def set_paragraph_format(p, first_line=True, align=None, before=0, after=0):
    pf = p.paragraph_format
    pf.line_spacing = 1.5
    pf.space_before = Pt(before)
    pf.space_after = Pt(after)
    if first_line:
        pf.first_line_indent = Cm(0.74)
    if align is not None:
        p.alignment = align


def add_para(doc, text="", size=12, bold=False, first_line=True, align=None, before=0, after=0):
    p = doc.add_paragraph()
    set_paragraph_format(p, first_line=first_line, align=align, before=before, after=after)
    run = p.add_run(text)
    set_run_font(run, size=size, bold=bold)
    return p


def add_heading(doc, text, level=1):
    p = doc.add_paragraph()
    set_paragraph_format(p, first_line=False, before=0, after=6)
    run = p.add_run(text)
    set_run_font(run, size=12, bold=True)
    return p


def add_cover_line(doc, label, value, right_label=None, right_value=None):
    p = doc.add_paragraph()
    set_paragraph_format(p, first_line=False, align=WD_ALIGN_PARAGRAPH.CENTER, before=0, after=0)
    text = f"{label}{value}"
    if right_label:
        text += f"    {right_label}{right_value}"
    run = p.add_run(text)
    set_run_font(run, size=12)
    return p


def set_cell_text(cell, text, size=12, bold=False, align=WD_ALIGN_PARAGRAPH.CENTER):
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    cell.text = ""
    p = cell.paragraphs[0]
    set_paragraph_format(p, first_line=False, align=align, before=0, after=0)
    run = p.add_run(text)
    set_run_font(run, size=size, bold=bold)


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_table_borders(table):
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = OxmlElement(f"w:{edge}")
        tag.set(qn("w:val"), "single")
        tag.set(qn("w:sz"), "8")
        tag.set(qn("w:space"), "0")
        tag.set(qn("w:color"), "000000")
        borders.append(tag)
    tbl_pr.append(borders)


def add_caption(doc, text):
    p = doc.add_paragraph()
    set_paragraph_format(p, first_line=False, align=WD_ALIGN_PARAGRAPH.CENTER, before=0, after=3)
    run = p.add_run(text)
    set_run_font(run, size=10.5)
    return p


def add_table(doc, caption, headers, rows, widths=None):
    add_caption(doc, caption)
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    set_table_borders(table)
    for i, header in enumerate(headers):
        set_cell_text(table.rows[0].cells[i], header, size=10.5, bold=True)
        set_cell_shading(table.rows[0].cells[i], "D9EAF7")
        if widths:
            table.rows[0].cells[i].width = Cm(widths[i])
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            align = WD_ALIGN_PARAGRAPH.LEFT if i == len(row) - 1 or len(str(value)) > 12 else WD_ALIGN_PARAGRAPH.CENTER
            set_cell_text(cells[i], str(value), size=10.5, align=align)
            if widths:
                cells[i].width = Cm(widths[i])
    add_para(doc, "", size=12, first_line=False, after=0)
    return table


def add_section_title(doc):
    p = doc.add_paragraph()
    set_paragraph_format(p, first_line=False, align=WD_ALIGN_PARAGRAPH.CENTER, before=0, after=12)
    run = p.add_run("毕 业 设 计（论 文）开 题 报 告")
    set_run_font(run, size=14, bold=True)


def new_report_item(doc):
    doc.add_page_break()
    add_section_title(doc)


def setup_document(doc):
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(3.0)
    section.right_margin = Cm(2.5)
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Times New Roman"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Times New Roman")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Times New Roman")
    normal.font.size = Pt(12)
    normal.paragraph_format.line_spacing = 1.5


def build_doc():
    doc = Document()
    setup_document(doc)

    # Cover page, following the supplied template's school-form style.
    add_para(doc, "青岛理工大学信息与控制工程学院", size=14, bold=True, first_line=False, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_para(doc, "毕业设计(论文)开题报告", size=18, bold=True, first_line=False, align=WD_ALIGN_PARAGRAPH.CENTER, after=24)
    for _ in range(3):
        add_para(doc, "", first_line=False)
    add_cover_line(doc, "学 生 姓 名：金宁丰", "", "学 号：", "202201050631")
    add_cover_line(doc, "班级：软件222", "")
    add_cover_line(doc, "设计(论文)题目：智能考研知识库系统的设计与实现——基于混合检索与RAG技术", "")
    add_cover_line(doc, "指 导 教 师：陈晓琳", "")
    add_cover_line(doc, "教   研   室：软件工程", "")
    for _ in range(8):
        add_para(doc, "", first_line=False)
    add_para(doc, "2026年 4 月 15 日", size=12, first_line=False, align=WD_ALIGN_PARAGRAPH.CENTER)

    new_report_item(doc)
    add_heading(doc, "1．本课题的研究背景、意义")
    add_heading(doc, "1.1 研究背景")
    paragraphs = [
        "随着生成式人工智能、大语言模型和教育信息化技术的发展，在线学习系统正在从单纯的资料存储与课程展示，逐步转向面向学习者个体需求的知识组织、智能检索和过程辅助。考研学习具有周期长、资料类型多、知识点关联强和复盘频率高等特点，学习者通常需要同时管理教材、讲义、PDF、Word 笔记、图片资料、历年真题和错题记录。如果缺少统一的知识管理工具，资料会分散在本地文件夹、网盘、聊天记录和纸质笔记中，导致查找成本高、复习线索断裂，难以形成稳定的备考闭环。",
        "传统学习平台多采用“文件存储 + 关键词搜索”的交互方式，只能在文件名或文本中寻找表层匹配，难以理解用户自然语言问题中的语义意图。例如，用户提出综合性、概念性或跨章节问题时，关键词检索往往不能准确召回相关片段；当问题与资料中的表述不完全一致时，系统也难以识别同义表达和上下文关系。与此同时，普通大模型问答虽然具备较强的语言生成能力，但如果不能接入用户私有资料，就容易出现回答泛化、依据不清和来源不可追溯等问题。",
        "检索增强生成（Retrieval-Augmented Generation，RAG）技术为上述问题提供了新的解决思路。RAG 通过在大语言模型生成答案之前引入外部知识检索过程，将用户资料中的相关片段作为上下文输入模型，使答案能够建立在显式知识来源之上。结合关键词检索、向量检索、排序融合和来源引用机制，系统既可以保留传统检索对精确术语的匹配能力，又可以利用语义向量弥补自然语言表达差异带来的召回不足。因此，将 RAG 技术引入考研学习场景，能够在资料管理、知识问答、错题复盘和学习规划之间建立更紧密的联系。",
        "本课题面向个人考研学习需求，设计并实现“智能考研知识库系统”。系统采用 Vue3 + Spring Boot 3 的前后端分离架构，围绕用户认证、资料上传、文本抽取、OCR 识别、知识切片、混合检索、RAG 问答、学习计划、错题集和 AI 设置等模块展开设计，目标是将分散资料转化为可检索、可引用、可辅助决策的个人知识库。"
    ]
    for text in paragraphs:
        add_para(doc, text)
    add_heading(doc, "1.2 研究意义")
    for text in [
        "从理论意义看，本课题将检索增强生成、混合检索和个性化学习支持思想结合到考研自学习场景中，探索“资料上传--知识切片--混合召回--上下文增强--答案生成--来源核验”的系统化路径。该路径有助于说明大语言模型在垂直学习场景中不能只作为通用聊天工具使用，而应与用户私有资料、检索机制和业务流程协同设计。",
        "从技术意义看，本课题覆盖前后端分离、JWT 鉴权、文件抽取、OCR、知识片段管理、Elasticsearch 检索、Embedding 向量召回、RRF 融合排序、SSE 流式问答和 OpenAI 兼容接口调用等关键技术。通过对这些技术进行工程化组合，可以形成一个结构清晰、边界明确、可扩展性较好的智能学习系统原型。",
        "从实践意义看，系统能够帮助考研学习者统一管理资料、快速定位知识点、基于私有资料进行答疑、记录错题并进行随机练习，还可以通过 AI 生成学习计划草稿。相较于单一文件管理工具或普通聊天机器人，本系统更强调学习过程中的资料、问题、计划和复盘之间的闭环关系，有助于提高复习效率和学习自我管理能力。"
    ]:
        add_para(doc, text)

    new_report_item(doc)
    add_heading(doc, "2．课题的主要设计（论文）内容")
    add_heading(doc, "2.1 系统总体架构设计")
    for text in [
        "本课题拟构建一套面向个人考研学习者的智能知识库系统，采用前后端分离和后端分层设计。前端基于 Vue3、Vite 和 Axios 实现单页应用，提供登录注册、我的知识库、知识问答、学习规划、错题集和 AI 设置等界面；后端基于 Spring Boot 3、Spring MVC、Spring Security、Spring Data JPA 和 JWT 实现业务接口、用户鉴权、资源归属校验和数据访问；数据层使用 H2 或 MySQL 保存用户、文件夹、资料、知识片段、学习计划、错题和 AI 参数，文件原件存储在本地 uploads 目录，Elasticsearch 作为可选检索增强服务。",
        "系统整体按照表现层、接口层、业务层、数据层和外部能力层进行组织。表现层负责页面渲染和交互；接口层提供 RESTful API、SSE 流式问答接口和 JWT 鉴权入口；业务层划分为认证、文件夹、资料文件、文本抽取、知识问答、学习计划、错题管理和 AI 设置等服务；数据层负责结构化数据、文件和知识片段的存储；外部能力层包括 Tesseract OCR、Elasticsearch、Embedding 模型和 OpenAI 兼容聊天模型。"
    ]:
        add_para(doc, text)
    add_heading(doc, "2.2 核心功能模块设计")
    for text in [
        "（1）用户认证与数据隔离模块。实现用户注册、登录、密码哈希存储和 JWT 签发，除注册登录外的业务接口均需要携带 Authorization 请求头。后端在访问文件夹、文件、错题和学习计划时校验当前用户归属，保证不同用户的数据相互隔离。",
        "（2）资料管理与知识库构建模块。用户可以创建多级资料文件夹，上传 PDF、Word、图片、文本和 Markdown 等资料。系统根据文件类型调用 PDFBox、Apache POI、Tesseract OCR 或文本读取逻辑抽取内容，允许用户手动编辑抽取文本，并将资料加入或移出知识库。加入知识库后，系统按固定窗口和重叠策略切分为 KnowledgeChunk，为后续检索问答提供基础。",
        "（3）基于 RAG 的知识问答模块。用户选择文件夹后，可在该范围及其子文件夹内提问。系统先执行混合检索，召回与问题相关的知识片段，再进行规则重排序和多样性筛选，最后构造包含资料片段和引用要求的 Prompt 调用大模型生成答案。系统支持答疑助手、定制练题、深度回答、来源引用和 SSE 流式输出；当 Elasticsearch、Embedding 或模型接口不可用时，系统提供本地关键词检索和摘要兜底。",
        "（4）学习计划智能辅助模块。系统支持按周查看计划，手动新增、修改和删除课程、自习、复盘、考试、任务和休息等类型的计划项。AI 规划功能通过对话收集用户学习目标与时间约束，生成计划操作草稿，用户确认后再写入真实日程，从而避免模型直接修改用户数据带来的风险。",
        "（5）错题集与复盘模块。用户可以录入错题文本、上传题目或解析附件，通过 OCR 或文件识别补充文本内容；系统支持掌握状态管理、科目标签分类、筛选浏览和随机练习。该模块用于将零散错题结构化，帮助用户围绕薄弱知识点进行持续复盘。",
        "（6）AI 设置与模型配置模块。用户可以配置聊天模型、Embedding 模型、API Endpoint、API Key、系统 Prompt 和 AI 角色定位，使系统能够兼容 OpenAI、DeepSeek、通义千问等 OpenAI 兼容接口，也便于在无 API Key 时退化为本地检索演示。"
    ]:
        add_para(doc, text)
    add_heading(doc, "2.3 关键技术与创新点")
    for text in [
        "本课题的关键技术包括：基于 Spring Security 和 JWT 的身份认证；基于 Apache POI、PDFBox 和 Tesseract 的多类型资料文本抽取；基于 KnowledgeChunk 的知识切片与本地存储；基于 Elasticsearch multi_match 和 dense_vector 的关键词检索与向量检索；基于 RRF 的多路召回结果融合；基于 Prompt Engineering 的知识增强问答；基于 SSE 的流式回答；以及面向外部服务不可用场景的本地降级机制。",
        "系统创新点主要体现在三个方面：一是将考研资料管理与 RAG 问答结合，实现面向个人私有资料的来源可追溯问答；二是采用关键词检索、向量检索、RRF 融合、规则 rerank 和多样性选择构成多阶段检索优化流程；三是将 AI 能力嵌入知识库、学习计划和错题复盘业务，使系统不只停留在聊天层面，而是服务于完整的考研学习闭环。"
    ]:
        add_para(doc, text)

    new_report_item(doc)
    add_heading(doc, "3．研究方案及进度计划")
    add_heading(doc, "3.1 研究目标")
    for text in [
        "本课题的总体目标是设计并实现一套“智能考研知识库系统”，使用户能够将分散学习资料统一上传、抽取、切片并加入知识库，在此基础上通过混合检索与 RAG 技术完成智能问答，同时结合学习计划和错题复盘模块形成自学习辅助闭环。",
        "具体目标包括：完成前后端分离系统架构设计；实现用户认证与数据隔离；实现资料上传、文本抽取、OCR 和知识库管理；实现混合检索、RAG 问答、来源引用和流式输出；实现学习计划与 AI 规划草稿；实现错题录入、标签状态管理和随机练习；完成系统测试、用户手册、毕业论文和答辩材料。"
    ]:
        add_para(doc, text)
    add_heading(doc, "3.2 技术路线")
    add_table(
        doc,
        "表3-1 技术栈表",
        ["层级", "技术组件", "功能定位"],
        [
            ["前端表现层", "Vue3、Vite、Axios", "用户界面渲染、页面状态管理、REST/SSE 请求"],
            ["后端业务层", "Java 21、Spring Boot 3、Spring MVC", "业务接口、服务编排、异常处理"],
            ["安全认证层", "Spring Security、JWT", "登录认证、接口保护、用户数据隔离"],
            ["数据持久层", "Spring Data JPA、Hibernate、H2/MySQL", "用户、资料、计划、错题和知识片段存储"],
            ["文本处理层", "PDFBox、Apache POI、Tesseract OCR", "PDF、Word、图片和文本资料抽取"],
            ["检索增强层", "Elasticsearch、Embedding、RRF", "关键词检索、向量检索、融合排序和召回优化"],
            ["AI 能力层", "OpenAI Compatible API、Prompt Engineering、SSE", "知识增强问答、学习计划草稿生成和流式回答"],
        ],
        widths=[3.2, 5.0, 7.0],
    )
    add_heading(doc, "3.3 研究方法")
    for text in [
        "（1）文献研究法。查阅个性化学习路径推荐、教育知识图谱、大语言模型、RAG、向量检索、排序融合和教育智能问答等相关研究，明确本课题的理论基础和技术边界。",
        "（2）需求分析法。结合考研学习场景，分析资料管理、知识检索、问答来源、学习计划和错题复盘等需求，形成软件需求说明书和功能模块划分。",
        "（3）原型迭代法。按照软件工程流程进行概要设计、详细设计、编码实现、联调测试和文档完善，通过可运行系统持续验证设计方案。",
        "（4）实验验证法。围绕资料上传、文本抽取、知识库切片、混合检索、问答生成、来源引用、计划管理和错题练习等功能设计测试用例，验证系统正确性与稳定性。"
    ]:
        add_para(doc, text)
    add_heading(doc, "3.4 拟解决的关键问题")
    for text in [
        "（1）多类型学习资料的文本抽取问题。针对 PDF、Word、图片和普通文本采用不同抽取策略，并在抽取失败时允许用户手动编辑和补全文本，保证知识库构建的基础质量。",
        "（2）私有知识库问答的检索质量问题。通过关键词检索和向量检索互补召回，再使用 RRF 融合、规则 rerank 和多样性筛选减少无关片段进入 Prompt 的概率，提高回答依据的相关性。",
        "（3）AI 回答的可信度与可追溯问题。系统在 Prompt 中约束模型优先依据知识库内容回答，并在答案中返回来源编号、文件名、页码和片段摘要，帮助用户核验原始资料。",
        "（4）外部服务不可用的系统稳定性问题。针对 Elasticsearch、Embedding 和大模型 API 不可用的情况设计本地检索与摘要兜底，保证毕业设计演示和核心流程不完全依赖外部服务。",
        "（5）AI 与学习业务融合问题。学习计划模块采用“AI 生成草稿--用户确认--写入日程”的机制，错题模块则通过状态、标签和随机练习支持复盘，避免 AI 功能停留在孤立聊天层面。"
    ]:
        add_para(doc, text)
    add_heading(doc, "3.5 进度计划")
    add_table(
        doc,
        "表3-2 进度计划表",
        ["阶段", "起止日期", "主要工作内容"],
        [
            ["第一阶段", "2026年04月15日--2026年04月26日", "系统功能完善与联调测试，完成 AI 模块与业务系统集成，修复核心缺陷"],
            ["第二阶段", "2026年04月27日--2026年05月11日", "编写用户手册，整理系统设计文档，完成毕业设计论文初稿"],
            ["第三阶段", "2026年05月12日--2026年05月18日", "根据指导教师意见修改论文，补充测试内容，完善格式与参考文献"],
            ["第四阶段", "2026年05月19日--2026年05月30日", "整理各类文档，根据评阅意见修改论文，完善系统功能并准备答辩"],
            ["第五阶段", "2026年06月01日--2026年06月06日", "完成论文查重检测与最终修改，提交正式论文及相关材料"],
            ["第六阶段", "2026年06月07日--2026年06月11日", "参加毕业答辩，根据答辩意见修改论文，完成毕业答辩材料归档"],
        ],
        widths=[2.5, 5.0, 7.7],
    )

    new_report_item(doc)
    add_heading(doc, "4．参考文献")
    refs = [
        "云岳, 代欢, 张育培, 等. 个性化学习路径推荐综述[J]. 软件学报, 2022, 33(12): 4590-4615.",
        "李惠乾, 钟柏昌. 教育知识图谱: 研究进展与未来发展--基于2013-2023年中文核心期刊载文的分析[J]. 计算机工程, 2024, 50(7): 1-12.",
        "NABIZADEH A H, GONCALVES D, GAMA S, et al. Learning path personalization and recommendation methods: a survey of the state-of-the-art[J]. Expert Systems with Applications, 2020, 159: 113596.",
        "HOLMES W, BIALIK M, FADEL C. Artificial intelligence in education: promises and implications for teaching and learning[M]. Boston: Center for Curriculum Redesign, 2019.",
        "LEWIS P, PEREZ E, PIKTUS A, et al. Retrieval-augmented generation for knowledge-intensive NLP tasks[C]//Advances in Neural Information Processing Systems 33. 2020: 9459-9474.",
        "KARPUKHIN V, OGUZ B, MIN S, et al. Dense passage retrieval for open-domain question answering[C]//Proceedings of the 2020 Conference on Empirical Methods in Natural Language Processing. Stroudsburg: Association for Computational Linguistics, 2020: 6769-6781.",
        "CORMACK G V, CLARKE C L A, BUTTCHER S. Reciprocal rank fusion outperforms Condorcet and individual rank learning methods[C]//Proceedings of the 32nd International ACM SIGIR Conference on Research and Development in Information Retrieval. New York: ACM, 2009: 758-759.",
        "MALKOV Y A, YASHUNIN D A. Efficient and robust approximate nearest neighbor search using hierarchical navigable small world graphs[J]. IEEE Transactions on Pattern Analysis and Machine Intelligence, 2020, 42(4): 824-836.",
        "BROWN T B, MANN B, RYDER N, et al. Language models are few-shot learners[C]//Advances in Neural Information Processing Systems 33. 2020: 1877-1901.",
        "WEI J, WANG X, SCHUURMANS D, et al. Chain-of-thought prompting elicits reasoning in large language models[C]//Advances in Neural Information Processing Systems 35. 2022: 24824-24837.",
        "GAO Y, XIONG Y, GAO X, et al. Retrieval-augmented generation for large language models: a survey[EB/OL]. (2023-12-18)[2026-05-18]. https://arxiv.org/abs/2312.10997.",
        "CHEN J, LIN H, HAN X, SUN L. Benchmarking large language models in retrieval-augmented generation[EB/OL]. (2023-09-04)[2026-05-18]. https://arxiv.org/abs/2309.01431.",
        "刘雪颖, 云静, 李博, 等. 基于大型语言模型的检索增强生成综述[J]. 计算机工程与应用, 2025, 61(13): 1-25.",
        "曹荣荣, 柳林, 于艳东, 等. 融合知识图谱的大语言模型研究综述[J]. 计算机应用研究, 2025, 42(8): 2255-2266.",
        "NG C, FUNG Y. Educational personalized learning path planning with large language models[EB/OL]. (2024-07-16)[2026-05-18]. https://arxiv.org/abs/2407.11773.",
        "LI Z, WANG Z, WANG W, et al. Retrieval-augmented generation for educational application: a systematic survey[J]. Computers and Education: Artificial Intelligence, 2025, 8: 100417.",
        "SWACHA J, GRACEL M. Retrieval-Augmented Generation (RAG) chatbots for education: a survey of applications[J]. Applied Sciences, 2025, 15(8): 4234.",
        "YAN S Q, GU J C, ZHU Y, et al. Corrective retrieval augmented generation[EB/OL]. arXiv:2401.15884, 2024.",
        "YU Y, PING W, LIU Z, et al. RankRAG: unifying context ranking with retrieval-augmented generation in LLMs[C]//Advances in Neural Information Processing Systems 37. 2024.",
        "ZHAO Q, WANG R, CEN Y, et al. LongRAG: a dual-perspective retrieval-augmented generation paradigm for long-context question answering[C]//Proceedings of the 2024 Conference on Empirical Methods in Natural Language Processing. Stroudsburg: Association for Computational Linguistics, 2024: 22600-22632.",
        "ASAI A, WU Z, WANG Y, et al. Self-RAG: learning to retrieve, generate, and critique through self-reflection[C]//The Twelfth International Conference on Learning Representations. 2024.",
        "XU H, GAN W, QI Z, et al. Large language models for education: a survey[EB/OL]. arXiv:2405.13001, 2024.",
        "李晓理, 刘春芳, 耿劭坤. 知识图谱与大语言模型协同共生模式及其教育应用综述[J]. 计算机工程与应用, 2025, 61(15): 1-13.",
        "ABU-SALIH B, ALOTAIBI S. A systematic literature review of knowledge graph construction and application in education[J]. Heliyon, 2024, 10(3): e25383.",
        "陈恒. SSM+Spring Boot+Vue.js 3全栈开发从入门到实战[M]. 北京: 清华大学出版社, 2022.",
        "周红亮. Spring Boot 3核心技术与最佳实践[M]. 北京: 电子工业出版社, 2023.",
        "罗刚, 张子宪. Elasticsearch搜索引擎开发实战[M]. 北京: 机械工业出版社, 2018.",
        "YU Y, EVAN Y. Vue.js Guide[EB/OL]. (2024)[2026-05-18]. https://vuejs.org/guide/introduction.html.",
        "ELASTIC. Elasticsearch Guide[EB/OL]. (2024)[2026-05-18]. https://www.elastic.co/guide/.",
        "SPRING.IO. Spring Boot Reference Documentation[EB/OL]. (2024)[2026-05-18]. https://docs.spring.io/spring-boot/.",
    ]
    for i, ref in enumerate(refs, 1):
        p = add_para(doc, f"[{i}] {ref}", size=10.5, first_line=False)
        p.paragraph_format.left_indent = Cm(0.74)
        p.paragraph_format.first_line_indent = Cm(-0.74)

    doc.add_page_break()
    add_section_title(doc)
    add_heading(doc, "指导教师意见：")
    for _ in range(8):
        add_para(doc, "", first_line=False)
    add_para(doc, "指导教师：                         年    月    日", first_line=False, align=WD_ALIGN_PARAGRAPH.RIGHT)
    add_heading(doc, "审查小组意见：")
    for _ in range(8):
        add_para(doc, "", first_line=False)
    add_para(doc, "负责人：                           年    月    日", first_line=False, align=WD_ALIGN_PARAGRAPH.RIGHT)
    add_para(doc, "成 绩：", first_line=False)

    doc.save(OUT)
    return OUT


if __name__ == "__main__":
    print(build_doc())
