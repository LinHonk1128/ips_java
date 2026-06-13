from pathlib import Path

from docx import Document
from docx.enum.text import WD_BREAK
from docx.shared import RGBColor


ROOT = Path(__file__).resolve().parents[1]
SOURCE = Path(r"C:\Users\admin\Desktop\毕业设计\论文改稿前.docx")
CHAPTER = ROOT / "docs" / "chapter4-system-design-draft.md"
OUTPUT = ROOT / "docs" / "论文改稿前-含第4章系统设计.docx"


def add_markdown_line(doc: Document, line: str) -> None:
    text = line.strip()
    if not text:
        return

    if text.startswith("# "):
        doc.add_heading(text[2:].strip(), level=1)
        return
    if text.startswith("## "):
        doc.add_heading(text[3:].strip(), level=2)
        return
    if text.startswith("### "):
        doc.add_heading(text[4:].strip(), level=3)
        return

    paragraph = doc.add_paragraph()
    run = paragraph.add_run(text)
    if text.startswith("【此处插入图"):
        run.bold = True
        run.font.color.rgb = RGBColor(192, 0, 0)


def main() -> None:
    doc = Document(str(SOURCE))

    if len(doc.paragraphs) > 0:
        doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)

    for raw_line in CHAPTER.read_text(encoding="utf-8").splitlines():
        add_markdown_line(doc, raw_line)

    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(OUTPUT))
    print(OUTPUT)


if __name__ == "__main__":
    main()
