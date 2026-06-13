import argparse
import json
import zipfile
from pathlib import Path

from docx import Document


def paragraph_text(paragraph):
    return "".join(run.text for run in paragraph.runs)


def inspect_docx(path):
    doc = Document(path)
    result = {
        "path": str(path),
        "paragraphs": [],
        "tables": [],
        "sections": [],
        "inline_shapes": len(doc.inline_shapes),
    }

    for index, paragraph in enumerate(doc.paragraphs):
        text = paragraph_text(paragraph)
        if text.strip() or paragraph.style.name:
            result["paragraphs"].append(
                {
                    "index": index,
                    "style": paragraph.style.name,
                    "text": text,
                    "runs": [
                        {
                            "text": run.text,
                            "bold": run.bold,
                            "italic": run.italic,
                            "size_pt": run.font.size.pt if run.font.size else None,
                            "font": run.font.name,
                        }
                        for run in paragraph.runs
                    ],
                }
            )

    for table_index, table in enumerate(doc.tables):
        rows = []
        for row_index, row in enumerate(table.rows):
            cells = []
            for cell_index, cell in enumerate(row.cells):
                cells.append(
                    {
                        "index": cell_index,
                        "text": "\n".join(
                            paragraph_text(paragraph) for paragraph in cell.paragraphs
                        ),
                        "paragraphs": [
                            {
                                "style": paragraph.style.name,
                                "text": paragraph_text(paragraph),
                                "runs": [
                                    {
                                        "text": run.text,
                                        "bold": run.bold,
                                        "italic": run.italic,
                                        "size_pt": (
                                            run.font.size.pt if run.font.size else None
                                        ),
                                        "font": run.font.name,
                                    }
                                    for run in paragraph.runs
                                ],
                            }
                            for paragraph in cell.paragraphs
                        ],
                    }
                )
            rows.append({"index": row_index, "cells": cells})
        result["tables"].append(
            {
                "index": table_index,
                "rows": rows,
                "style": table.style.name if table.style else None,
            }
        )

    for index, section in enumerate(doc.sections):
        result["sections"].append(
            {
                "index": index,
                "width": section.page_width,
                "height": section.page_height,
                "top_margin": section.top_margin,
                "bottom_margin": section.bottom_margin,
                "left_margin": section.left_margin,
                "right_margin": section.right_margin,
                "header_text": "\n".join(
                    paragraph_text(paragraph)
                    for paragraph in section.header.paragraphs
                ),
                "footer_text": "\n".join(
                    paragraph_text(paragraph)
                    for paragraph in section.footer.paragraphs
                ),
            }
        )

    with zipfile.ZipFile(path) as archive:
        result["package_parts"] = sorted(archive.namelist())

    return result


def main():
    parser = argparse.ArgumentParser()
    parser.add_argument("paths", nargs="+", type=Path)
    parser.add_argument("--output-dir", type=Path, required=True)
    args = parser.parse_args()
    args.output_dir.mkdir(parents=True, exist_ok=True)

    for path in args.paths:
        result = inspect_docx(path)
        output = args.output_dir / f"{path.stem}.json"
        output.write_text(
            json.dumps(result, ensure_ascii=False, indent=2), encoding="utf-8"
        )
        print(output)


if __name__ == "__main__":
    main()
