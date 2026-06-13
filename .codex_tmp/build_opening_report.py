from copy import deepcopy
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt


TEMPLATE = Path(".codex_tmp/template.docx")
OUTPUT = Path("开题报告-基于RAG与知识画像的考研学习辅助系统-金宁丰.docx")

FONT_CN = "宋体"
FONT_EN = "Times New Roman"
BODY_SIZE = Pt(12)
REF_SIZE = Pt(10.5)
LINE_SPACING = Pt(20)


def set_run_font(run, size=BODY_SIZE, bold=False, underline=None):
    run.font.name = FONT_EN
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), FONT_CN)
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), FONT_EN)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), FONT_EN)
    run.font.size = size
    run.bold = bold
    if underline is not None:
        run.underline = underline


def clear_paragraph(paragraph):
    for child in list(paragraph._p):
        if child.tag != qn("w:pPr"):
            paragraph._p.remove(child)


def set_cover_field(paragraph, label, value):
    clear_paragraph(paragraph)
    label_run = paragraph.add_run(label)
    set_run_font(label_run, Pt(15))
    value_run = paragraph.add_run(value if value else " " * 24)
    set_run_font(value_run, Pt(15), underline=True)


def configure_paragraph(paragraph, first_line=True, align=WD_ALIGN_PARAGRAPH.JUSTIFY):
    paragraph.alignment = align
    fmt = paragraph.paragraph_format
    fmt.space_before = Pt(0)
    fmt.space_after = Pt(0)
    fmt.line_spacing = LINE_SPACING
    fmt.first_line_indent = Pt(24) if first_line else Pt(0)
    fmt.widow_control = True


def set_cell_margins(cell, top=100, start=120, bottom=100, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for name, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{name}"))
        if node is None:
            node = OxmlElement(f"w:{name}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def reset_cell(cell):
    tc_pr = cell._tc.tcPr
    tc_pr_copy = deepcopy(tc_pr) if tc_pr is not None else None
    for child in list(cell._tc):
        cell._tc.remove(child)
    if tc_pr_copy is not None:
        cell._tc.append(tc_pr_copy)
    paragraph = OxmlElement("w:p")
    cell._tc.append(paragraph)
    set_cell_margins(cell)
    return cell.paragraphs[0]


def add_heading(cell, text, use_existing=False):
    paragraph = cell.paragraphs[-1] if use_existing else cell.add_paragraph()
    clear_paragraph(paragraph)
    configure_paragraph(paragraph, first_line=False, align=WD_ALIGN_PARAGRAPH.LEFT)
    paragraph.paragraph_format.keep_with_next = True
    run = paragraph.add_run(text)
    set_run_font(run, BODY_SIZE, bold=True)
    return paragraph


def add_body(cell, text, first_line=True):
    paragraph = cell.add_paragraph()
    configure_paragraph(paragraph, first_line=first_line)
    run = paragraph.add_run(text)
    set_run_font(run)
    return paragraph


def add_reference(cell, text):
    paragraph = cell.add_paragraph()
    configure_paragraph(paragraph, first_line=False, align=WD_ALIGN_PARAGRAPH.LEFT)
    paragraph.paragraph_format.line_spacing = Pt(16)
    run = paragraph.add_run(text)
    set_run_font(run, REF_SIZE)
    return paragraph


def add_blank(cell, count=1, line_spacing=Pt(20)):
    for _ in range(count):
        paragraph = cell.add_paragraph()
        configure_paragraph(paragraph, first_line=False)
        paragraph.paragraph_format.line_spacing = line_spacing


doc = Document(TEMPLATE)

set_cover_field(doc.paragraphs[11], "学    院 ", "信息管理学院")
set_cover_field(doc.paragraphs[12], "专业班级 ", "软件222")
set_cover_field(doc.paragraphs[13], "学生学号 ", "202201050631")
set_cover_field(doc.paragraphs[14], "学生姓名 ", "金宁丰")
set_cover_field(doc.paragraphs[15], "指导教师 ", "陈晓琳")
set_cover_field(doc.paragraphs[16], "指导教师职称 ", "")

date_paragraph = doc.paragraphs[19]
clear_paragraph(date_paragraph)
date_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
date_run = date_paragraph.add_run("2026年 4 月 15 日")
set_run_font(date_run, Pt(14))

table = doc.tables[0]

title_cell = table.rows[1].cells[0]
title_paragraph = reset_cell(title_cell)
configure_paragraph(title_paragraph, first_line=False, align=WD_ALIGN_PARAGRAPH.CENTER)
title_run = title_paragraph.add_run("基于RAG与知识画像的考研学习辅助系统设计与实现")
set_run_font(title_run, BODY_SIZE)

source_cell = table.rows[3].cells[0]
source_paragraph = reset_cell(source_cell)
configure_paragraph(source_paragraph, first_line=False, align=WD_ALIGN_PARAGRAPH.CENTER)
source_run = source_paragraph.add_run("指导教师拟定课题")
set_run_font(source_run, BODY_SIZE)

purpose_cell = table.rows[4].cells[0]
reset_cell(purpose_cell)
add_heading(purpose_cell, "一、本课题的研究目的和意义", use_existing=True)
add_body(
    purpose_cell,
    "考研备考具有周期长、知识量大、资料来源多样等特点。学习者通常需要同时管理教材、讲义、个人笔记、历年真题、错题截图和专题总结，容易出现资料分散、知识定位效率低、错题复盘不连续、薄弱点难以识别以及学习计划缺少动态依据等问题。现有在线学习平台和电子笔记工具能够提供课程获取或资料存储功能，但对个人资料的语义理解、来源追溯和学习状态联动支持仍较有限。",
)
add_body(
    purpose_cell,
    "大语言模型具备较强的自然语言理解与生成能力，但直接用于学习问答时可能存在知识幻觉、回答依据不明确等问题。检索增强生成（RAG）通过在答案生成前检索外部知识内容，可使模型围绕指定资料进行回答并返回来源依据。若进一步将问答引用、用户反馈和错题练习结果关联到具体知识片段，还可形成面向个人学习过程的知识画像，为薄弱点识别和复习计划调整提供数据基础。",
)
add_body(
    purpose_cell,
    "本课题拟设计一套基于RAG与知识画像的考研学习辅助系统，将用户上传的多类型学习资料转化为可检索、可引用、可记录学习状态的知识片段；采用关键词检索与向量检索相结合的方式支持个人知识库问答；根据问答反馈、错题练习和知识片段事件构建知识画像，并据此辅助生成复习建议与学习计划，形成“资料整理—知识检索—错题复盘—薄弱点识别—计划调整”的学习闭环。",
)
add_body(
    purpose_cell,
    "本课题的研究意义主要体现在三个方面：在理论层面，探索RAG问答、学习行为分析与知识画像在个人考研场景中的结合方式；在实践层面，帮助学习者集中管理资料、快速定位知识依据并开展针对性复习；在工程层面，形成基于Spring Boot、Vue 3、MySQL、Elasticsearch和大语言模型接口的模块化实现方案，为智能学习系统和个人知识库平台的开发提供参考。",
)

status_cell = table.rows[5].cells[0]
reset_cell(status_cell)
add_heading(status_cell, "二、国内外研究现状及发展趋势", use_existing=True)
add_body(
    status_cell,
    "国内外在线学习系统的研究已由课程资源发布和学习记录保存，逐步转向个性化学习路径、学习过程支持和资源推荐。云岳等从学习目标、知识基础和学习状态等角度总结了个性化学习路径推荐方法[1]；Nabizadeh等对学习者特征、资源关系和推荐策略进行了系统梳理[2]。相关研究为个性化学习提供了方法基础，但多数应用依赖固定课程结构和教师组织，难以直接适应考研自学中资料来源分散、复习节奏差异大等特点。",
)
add_body(
    status_cell,
    "在智能问答领域，传统系统主要依赖规则、关键词匹配或结构化知识库。大语言模型提升了自然语言理解和答案生成能力，但其参数知识难以及时覆盖用户私有资料，且可能产生无法核验的内容。Lewis等提出的RAG框架将外部检索结果作为生成模型的上下文，为知识密集型问答提供了可追溯路径[3]。近年来，国内外研究进一步关注知识库构建、文档切分、混合召回、候选重排和回答引用等环节[4-6]。",
)
add_body(
    status_cell,
    "在检索方法方面，关键词检索适合匹配明确术语，向量检索能够发现表述不同但语义相近的内容。RRF等排序融合方法可减少不同检索器分数尺度不一致带来的影响[7]，因此混合检索、多阶段重排和来源引用已成为RAG系统的重要发展方向。现有研究较多关注召回率和回答准确性，对问答后的学习反馈如何持续沉淀为个人学习状态研究相对不足。",
)
add_body(
    status_cell,
    "教育知识图谱、学习分析和知识追踪研究为学习诊断提供了另一类思路。相关工作通过组织知识点、学习资源、答题记录和学习行为，对掌握状态进行估计并支持个性化推荐[8-11]。但个人考研资料规模较小、格式不统一，直接构建复杂知识图谱或采用需要大量交互数据的知识追踪模型，存在成本较高和冷启动明显等问题。",
)
add_body(
    status_cell,
    "综合来看，相关研究正呈现三个趋势：一是从通用模型问答转向基于私有知识库的可追溯问答；二是从单一检索转向关键词、向量召回与重排序结合的混合检索；三是从孤立的问答或错题功能转向学习行为、薄弱点分析和计划推荐联动。面向考研自学场景，仍需研究一种数据要求较低、能够连接个人资料、问答反馈、错题复盘和学习计划的轻量化知识画像方案。",
)
add_heading(status_cell, "三、本课题主要研究内容")
for text in [
    "（1）系统需求与总体架构。分析考研学习者在资料管理、知识问答、错题复盘和学习计划方面的需求，拟采用前后端分离架构，划分用户与学习档案、资料管理、知识库问答、知识画像、错题复盘和学习计划等模块。",
    "（2）个人知识库构建。研究PDF、Word、图片和文本等资料的文本抽取方法，结合标题、句子边界和语义连续性进行知识切片，将长篇资料转化为可检索、可引用、可记录学习事件的KnowledgeChunk。",
    "（3）混合检索与RAG问答。研究关键词检索与Embedding向量检索的互补召回机制，采用RRF进行结果融合，并结合规则重排、去重和多样性筛选构造模型上下文，使回答能够基于个人资料生成并提供来源引用。",
    "（4）知识画像分析。围绕知识片段汇总问答引用、用户反馈、错题关联、练习结果和复习时间等事件，设计掌握程度、错误频率、遗忘风险和复习优先级等指标，形成可动态更新的个人知识画像。",
    "（5）错题复盘与知识点关联。支持错题录入、附件识别、标签分类、状态管理和练习记录，将错题与相关知识片段建立关联，使练习结果能够回写并更新画像。",
    "（6）个性化学习计划与系统验证。根据知识画像和错题分析结果生成学习建议与计划草稿，经用户确认后写入日程；围绕核心业务流程、数据隔离、异常降级和系统易用性设计测试方案。",
]:
    add_body(status_cell, text, first_line=False)

method_cell = table.rows[6].cells[0]
reset_cell(method_cell)
add_heading(method_cell, "四、拟解决的关键问题", use_existing=True)
for text in [
    "（1）多类型资料难以统一处理的问题。不同文件格式在文本结构、编码和图片内容上存在差异，需要设计分类抽取、OCR识别、人工校正和语义切片流程，保证进入知识库的文本质量。",
    "（2）用户提问与资料原文表述不一致的问题。单一关键词检索可能遗漏语义相关内容，单一向量检索又可能弱化专业术语匹配，需研究混合召回、RRF融合和候选重排方法。",
    "（3）生成式回答可信度不足的问题。需要限制模型优先依据检索片段回答，并保留文件名、页码或片段摘要等来源信息，便于学习者核验。",
    "（4）个人学习数据稀疏条件下的画像构建问题。考研自学缺少标准化题库和大规模作答数据，需以知识片段为基本单元，利用问答反馈、错题练习和时间衰减构建轻量化、可解释的画像指标。",
    "（5）各功能模块数据难以形成闭环的问题。需明确资料、知识片段、问答事件、错题、画像和学习计划之间的数据关系，并设计外部检索或模型服务不可用时的降级策略。",
]:
    add_body(method_cell, text, first_line=False)
add_heading(method_cell, "五、研究思路和方法")
add_body(
    method_cell,
    "本课题拟按照“问题分析—需求建模—总体设计—模块实现—系统测试—总结改进”的思路开展。首先分析考研学习中资料分散、问答依据不清和复习反馈不足等问题，明确系统边界和核心数据对象；随后以知识片段作为连接资料内容与学习行为的中间载体，设计RAG问答和知识画像两条核心数据链路；最后通过原型迭代和测试验证各模块的可行性。",
)
for text in [
    "（1）文献研究法。查阅个性化学习、教育知识图谱、学习分析、知识追踪、RAG、混合检索和大语言模型应用等文献，梳理研究现状并确定技术边界。",
    "（2）需求分析与建模法。通过场景分析、用例分析和业务流程建模，明确用户角色、功能需求、非功能需求以及资料、知识片段、错题和计划等实体关系。",
    "（3）系统设计与原型迭代法。拟采用Spring Boot构建后端服务，Vue 3构建前端界面，MySQL存储业务数据，Elasticsearch提供全文和向量检索；按照模块逐步实现并通过联调持续修正设计。",
    "（4）对比与实验验证法。针对关键词检索、向量检索和混合检索设计典型问题集，比较召回结果；针对资料处理、来源引用、画像更新、错题回写和计划生成设计功能测试与异常场景测试。",
]:
    add_body(method_cell, text, first_line=False)
add_body(
    method_cell,
    "总体技术路线为：学习资料上传与文本抽取→语义切片与向量化→关键词检索和向量检索并行召回→RRF融合与规则重排→构造RAG上下文并生成带引用的回答→记录问答反馈和错题练习事件→更新知识画像→生成复习建议与学习计划草稿。",
)

schedule_cell = table.rows[7].cells[0]
reset_cell(schedule_cell)
add_heading(schedule_cell, "六、本课题的进度安排", use_existing=True)
for text in [
    "2026年4月15日—4月20日：完成文献调研、需求分析和可行性分析，明确研究目标、系统范围及总体技术路线。",
    "2026年4月21日—4月27日：完成系统总体架构、功能模块、业务流程和数据库结构设计，形成概要设计与详细设计文档。",
    "2026年4月28日—5月8日：完成用户认证、学习资料管理、文本抽取、知识切片和个人知识库基础功能。",
    "2026年5月9日—5月18日：完成混合检索、RAG问答、来源引用、知识画像、错题复盘和学习计划等核心模块，并开展前后端联调。",
    "2026年5月19日—5月26日：设计并执行功能测试、异常测试和核心流程验证，根据测试结果修正系统问题。",
    "2026年5月27日—6月3日：整理系统设计图、数据库表和测试材料，完成毕业论文初稿及用户使用说明。",
    "2026年6月4日—6月11日：根据指导意见修改论文和系统，完成格式审查、查重、答辩准备及材料归档。",
]:
    add_body(schedule_cell, text, first_line=False)

references_cell = table.rows[8].cells[0]
reset_cell(references_cell)
add_heading(references_cell, "七、参考文献", use_existing=True)
references = [
    "[1] 云岳, 代欢, 张育培, 等. 个性化学习路径推荐综述[J]. 软件学报, 2022, 33(12): 4590-4615.",
    "[2] NABIZADEH A H, GONCALVES D, GAMA S, et al. Learning path personalization and recommendation methods: a survey of the state-of-the-art[J]. Expert Systems with Applications, 2020, 159: 113596.",
    "[3] LEWIS P, PEREZ E, PIKTUS A, et al. Retrieval-augmented generation for knowledge-intensive NLP tasks[C]//Advances in Neural Information Processing Systems 33. 2020: 9459-9474.",
    "[4] 文森, 钱力, 胡懋地, 等. 基于大语言模型的问答技术研究进展综述[J]. 数据分析与知识发现, 2024, 8(6): 16-29.",
    "[5] 刘雪颖, 云静, 李博, 等. 基于大型语言模型的检索增强生成综述[J]. 计算机工程与应用, 2025, 61(13): 1-25.",
    "[6] SWACHA J, GRACEL M. Retrieval-Augmented Generation chatbots for education: a survey of applications[J]. Applied Sciences, 2025, 15(8): 4234.",
    "[7] CORMACK G V, CLARKE C L A, BÜTTCHER S. Reciprocal rank fusion outperforms Condorcet and individual rank learning methods[C]//Proceedings of the 32nd International ACM SIGIR Conference. 2009: 758-759.",
    "[8] 李惠乾, 钟柏昌. 教育知识图谱: 研究进展与未来发展——基于2013—2023年中文核心期刊载文的分析[J]. 计算机工程, 2024, 50(7): 1-12.",
    "[9] ALAMRI H A, WATSON S, WATSON W. A systematic review of the role of learning analytics in supporting personalized learning[J]. Education Sciences, 2024, 14(1): 51.",
    "[10] PIECH C, BASSEN J, HUANG J, et al. Deep knowledge tracing[C]//Advances in Neural Information Processing Systems 28. 2015.",
    "[11] 刘铁园, 陈威, 常亮, 等. 基于深度学习的知识追踪研究进展[J]. 计算机研究与发展, 2022, 59(1): 81-104.",
    "[12] 王士进, 吴金泽, 张浩天, 等. 可信的端到端深度学生知识画像建模方法[J]. 计算机研究与发展, 2023, 60(8): 1822-1833.",
    "[13] CARPENTER S K, PAN S C, BUTLER A C. The science of effective learning with spacing and retrieval practice[J]. Nature Reviews Psychology, 2022, 1: 496-511.",
    "[14] 车璐, 张志强, 周金佳, 等. 生成式人工智能的研究现状和发展趋势[J]. 科技导报, 2024, 42(12): 35-43.",
    "[15] 刘邦奇, 聂小林, 王士进, 等. 生成式人工智能与未来教育形态重塑：技术框架、能力特征及应用趋势[J]. 电化教育研究, 2024, 45(1): 13-20.",
]
for reference in references:
    add_reference(references_cell, reference)

add_blank(references_cell, 2)
add_heading(references_cell, "八、指导教师意见")
add_blank(references_cell, 5, line_spacing=Pt(24))
signature = references_cell.add_paragraph()
configure_paragraph(signature, first_line=False, align=WD_ALIGN_PARAGRAPH.RIGHT)
signature_run = signature.add_run("签名：                         ")
set_run_font(signature_run, BODY_SIZE)
signature_date = references_cell.add_paragraph()
configure_paragraph(signature_date, first_line=False, align=WD_ALIGN_PARAGRAPH.RIGHT)
signature_date_run = signature_date.add_run("年        月        日")
set_run_font(signature_date_run, BODY_SIZE)

doc.core_properties.title = "基于RAG与知识画像的考研学习辅助系统设计与实现开题报告"
doc.core_properties.subject = "本科毕业设计（论文）开题报告"
doc.core_properties.author = "金宁丰"
doc.core_properties.keywords = "RAG, 知识画像, 考研学习辅助系统"

doc.save(OUTPUT)
print(OUTPUT.resolve())
